#!/usr/bin/env python3
"""
Pitch Deck Word Document Generator — Thesis-branded .docx export.

Generates a professional Word document from any ticker's Pitch Deck pipeline output
with embedded chart images (verdict scorecard, financial trends, price ranges),
styled tables, citations, and Thesis branding.

Usage:
    python3 scripts/pdf/generate_pitch_deck_docx.py MNST
"""

import json
import os
import sys

# Ensure scripts/pdf is importable
sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', '..'))

from scripts.pdf.report_data_reader import ReportData
from scripts.pdf.safe_ticker import normalize_ticker
from scripts.pdf.section_renderers import (
    get_narrative, get_tables, get_red_flags, get_citations,
)
from scripts.pdf.chart_image_generator import (
    generate_bar_chart, generate_verdict_scorecard,
    generate_price_range_chart, generate_comparison_chart,
    generate_pipeline_flow, generate_donut, generate_radar,
    generate_stacked_bar,
)
from scripts.pdf.docx_helpers import (
    create_thesis_doc, add_title_page, add_styled_table,
    add_verdict_table, add_section_heading, add_body_paragraphs,
    embed_chart, add_red_flags as render_red_flags, add_citations_section,
    cleanup_temp_charts, VERDICT_COLORS_RGB, SLATE_600,
    render_verdict_box, render_accounting_red_flags, render_pre_decision_check,
)
from docx.shared import Pt


def _section_data_dict(section):
    if not isinstance(section, dict):
        return {}
    raw = section.get('data')
    if isinstance(raw, dict):
        return raw
    if isinstance(raw, str):
        try:
            parsed = json.loads(raw)
            return parsed if isinstance(parsed, dict) else {}
        except (ValueError, TypeError):
            return {}
    return {}


# =========================================================================
# PITCH DECK REDESIGN (2026-05-09) — KEY MIGRATION + GROUPING
# =========================================================================

# Legacy section-key → current section-key. Routes archived pitch-deck JSON
# (generated before the 2026-05-09 redesign) through the new top-level grouping.
LEGACY_KEY_MAP = {
    'radar': 'setup',
    'simple_predictable': 'business_quality',
    'simple_and_predictable': 'business_quality',
    'barriers_moats': 'moat_analysis',
    'fcf': 'cash_generation',
    'roe_roic_debt': 'returns_leverage',
    'management': 'management_capital_allocation',
    'pest': 'risk_profile',
    'pest_risks': 'risk_profile',
    'overall_verdict': 'investment_verdict',
    'valuation_summary': 'valuation',
}


def normalize_section_key(key):
    """Map legacy section keys to current keys; pass through if already current."""
    if not isinstance(key, str):
        return key
    return LEGACY_KEY_MAP.get(key, key)


# 8 top-level groupings, mirrored across PDF + DOCX + UI for consistency.
TOP_LEVEL_GROUPS = [
    {'title': 'Setup & Situation', 'subsection_keys': ['setup']},
    {'title': 'Business Quality', 'subsection_keys': ['business_quality']},
    {'title': 'Industry & Competitive Position',
     'subsection_keys': ['market_position', 'moat_analysis']},
    {'title': 'Financial Analysis',
     'subsection_keys': ['cash_generation', 'returns_leverage',
                         'balance_sheet', 'accounting_red_flags']},
    {'title': 'Management & Capital Allocation',
     'subsection_keys': ['management_capital_allocation']},
    {'title': 'Valuation', 'subsection_keys': ['valuation']},
    {'title': 'Risk Profile', 'subsection_keys': ['risk_profile']},
    {'title': 'Investment Verdict', 'subsection_keys': ['investment_verdict']},
]


# Chart treatments keyed by NORMALIZED (current) keys.
CHART_SECTIONS = {
    'business_quality': 'revenue',
    'cash_generation': 'fcf',
    'balance_sheet': 'debt_equity',
    'valuation': 'price_range',
}


def _add_revenue_chart(doc, data, temp_charts):
    """Add revenue trend chart for simple_and_predictable section."""
    try:
        years = data.get_financial_years(n=5)
        if not years:
            return
        rev_data = data.get_financial_data('income', 'revenues', years)
        if rev_data and len(rev_data) >= 2:
            chart_years = [str(y) for y, _ in rev_data]
            chart_values = [v / 1e9 for _, v in rev_data]
            path = generate_bar_chart(
                chart_years, chart_values,
                title='Revenue Trend (5yr)', unit='B',
            )
            temp_charts.append(path)
            embed_chart(doc, path)
    except Exception:
        pass


def _add_fcf_chart(doc, data, temp_charts):
    """Add FCF bar chart for fcf section."""
    try:
        years = data.get_financial_years(n=5)
        if not years:
            return
        ocf_data = data.get_financial_data('cashFlow', 'operating_cash_flow', years)
        capex_data = data.get_financial_data('cashFlow', 'capital_expenditures', years)
        if ocf_data and len(ocf_data) >= 2:
            # Compute FCF = OCF - CapEx
            ocf_map = dict(ocf_data)
            capex_map = dict(capex_data) if capex_data else {}
            fcf_years = []
            fcf_values = []
            for y in years:
                ocf = ocf_map.get(y)
                capex = abs(capex_map.get(y, 0))
                if ocf is not None:
                    fcf_years.append(str(y))
                    fcf_values.append((ocf - capex) / 1e9)
            if len(fcf_years) >= 2:
                path = generate_bar_chart(
                    fcf_years, fcf_values,
                    title='Free Cash Flow Trend', unit='B',
                )
                temp_charts.append(path)
                embed_chart(doc, path)
    except Exception:
        pass


def _add_debt_equity_chart(doc, data, temp_charts):
    """Add debt vs equity comparison chart for balance_sheet section."""
    try:
        years = data.get_financial_years(n=5)
        if not years:
            return
        debt_data = data.get_financial_data('balance', 'total_debt', years)
        equity_data = data.get_financial_data('balance', 'stockholders_equity', years)
        if debt_data and equity_data and len(debt_data) >= 2:
            chart_years = [str(y) for y, _ in debt_data]
            debt_vals = [v / 1e9 for _, v in debt_data]
            equity_vals = [v / 1e9 for _, v in equity_data[:len(debt_data)]]
            path = generate_comparison_chart(
                chart_years,
                [debt_vals, equity_vals],
                ['Total Debt', 'Equity'],
                title='Debt vs Equity', unit='B',
            )
            temp_charts.append(path)
            embed_chart(doc, path)
    except Exception:
        pass


def _add_moat_radar(doc, data, temp_charts):
    """Radar of moat strengths from moat_analysis section."""
    try:
        section = data.get_section('moat_analysis') or data.get_section('barriers_moats')
        raw = _section_data_dict(section)
        moats = raw.get('moatTypes')
        if not isinstance(moats, list) or len(moats) < 3:
            return
        strength_map = {'strong': 9, 'wide': 9, 'moderate': 6,
                        'narrow': 4, 'weak': 2, 'none': 0}
        labels, values = [], []
        for m in moats:
            if not isinstance(m, dict):
                continue
            t = str(m.get('type', '')).strip()
            s = str(m.get('strength', '')).lower().strip()
            if not t:
                continue
            labels.append(t.title())
            values.append(strength_map.get(s, 5))
        if len(labels) >= 3:
            path = generate_radar(labels, values, title='Moat Strength by Type', max_value=10)
            temp_charts.append(path)
            embed_chart(doc, path)
    except Exception:
        pass


def _add_capex_donut(doc, data, temp_charts):
    """Donut: maintenance vs growth capex from cash_generation section."""
    try:
        section = data.get_section('cash_generation') or data.get_section('fcf')
        raw = _section_data_dict(section)
        cb = raw.get('capexBreakdown')
        if not isinstance(cb, dict):
            return
        maint = cb.get('maintenanceCapex')
        growth = cb.get('growthCapex')
        if not (isinstance(maint, (int, float)) and isinstance(growth, (int, float))):
            return
        if maint + growth <= 0:
            return
        path = generate_donut([('Maintenance', maint), ('Growth', growth)],
                               title='Capex Split: Maintenance vs Growth')
        if path:
            temp_charts.append(path)
            embed_chart(doc, path)
    except Exception:
        pass


def _add_capital_allocation_stack(doc, data, temp_charts):
    """5-year stacked bar: CapEx + Buybacks + Dividends."""
    try:
        cf = data.data_packet.get('financials', {}).get('cashFlow', {})
        yrs = sorted([y for y in cf.keys() if str(y).isdigit()])[-5:]
        if len(yrs) < 3:
            return
        capex, buybacks, divs = [], [], []
        for y in yrs:
            rec = cf.get(y) or cf.get(str(y)) or {}
            capex.append(abs(rec.get('payments_to_acquire_property_plant_and_equipment') or
                             rec.get('capital_expenditures') or 0))
            buybacks.append(abs(rec.get('payments_for_repurchase_of_common_stock') or 0))
            divs.append(abs(rec.get('payments_of_dividends_common_stock') or
                            rec.get('payments_of_dividends') or 0))
        if not any(capex) and not any(buybacks) and not any(divs):
            return
        path = generate_stacked_bar(
            [str(y) for y in yrs],
            [capex, buybacks, divs],
            ['CapEx', 'Buybacks', 'Dividends'],
            title='Capital Allocation (5y)',
            unit='B',
        )
        temp_charts.append(path)
        embed_chart(doc, path)
    except Exception:
        pass


def _add_peer_comparison(doc, data, temp_charts):
    """Subject vs peer-median grouped bar for ROIC/ROE."""
    try:
        averages = data.data_packet.get('returnMetrics', {}).get('averages', {})
        subj_roic = averages.get('roic_3yr')
        subj_roe = averages.get('roe_3yr')
        if subj_roic is None or subj_roe is None:
            return
        peers = data.data_packet.get('peers') or []
        peer_roic, peer_roe = [], []
        for p in peers:
            if not isinstance(p, dict):
                continue
            rm = p.get('returnMetrics') or p.get('returns') or {}
            avgs = rm.get('averages') if isinstance(rm, dict) else None
            if isinstance(avgs, dict):
                if isinstance(avgs.get('roic_3yr'), (int, float)):
                    peer_roic.append(avgs['roic_3yr'])
                if isinstance(avgs.get('roe_3yr'), (int, float)):
                    peer_roe.append(avgs['roe_3yr'])

        def _median(xs):
            xs = sorted(xs)
            if not xs:
                return None
            n = len(xs)
            return xs[n // 2] if n % 2 else (xs[n // 2 - 1] + xs[n // 2]) / 2

        p_roic = _median(peer_roic)
        p_roe = _median(peer_roe)
        if p_roic is None and p_roe is None:
            return

        labels, subj_vals, peer_vals = [], [], []
        if subj_roic is not None and p_roic is not None:
            labels.append('ROIC (3y)')
            subj_vals.append(round(subj_roic, 1))
            peer_vals.append(round(p_roic, 1))
        if subj_roe is not None and p_roe is not None:
            labels.append('ROE (3y)')
            subj_vals.append(round(subj_roe, 1))
            peer_vals.append(round(p_roe, 1))
        if not labels:
            return
        path = generate_comparison_chart(
            labels, [subj_vals, peer_vals], ['Subject', 'Peer Median'],
            title='Peer Comparison — Subject vs. Peer Median', unit='%',
        )
        temp_charts.append(path)
        embed_chart(doc, path)
    except Exception:
        pass


def _add_price_range_chart(doc, data, temp_charts):
    """Add buy price range chart for valuation_summary section."""
    try:
        buy_prices = data.get_buy_prices()
        current_price = buy_prices.get('currentPrice', 0) or data.get_current_price()
        if not current_price:
            return

        methods = []
        # MOS
        mos = buy_prices.get('mosBuyPrice')
        if isinstance(mos, dict) and mos.get('low') and mos.get('high'):
            methods.append(('MOS', float(mos['low']), float(mos['high']), '#22c55e'))
        elif isinstance(mos, (int, float)) and mos > 0:
            methods.append(('MOS', float(mos) * 0.9, float(mos) * 1.1, '#22c55e'))

        # PBT
        pbt = buy_prices.get('pbtBuyPrice')
        if isinstance(pbt, dict) and pbt.get('low') and pbt.get('high'):
            methods.append(('PBT', float(pbt['low']), float(pbt['high']), '#3b82f6'))
        elif isinstance(pbt, (int, float)) and pbt > 0:
            methods.append(('PBT', float(pbt) * 0.9, float(pbt) * 1.1, '#3b82f6'))

        # Ten Cap
        tc = buy_prices.get('tenCapPrice')
        if isinstance(tc, dict):
            low = tc.get('ruleOne') or tc.get('low')
            high = tc.get('graham') or tc.get('high')
            if low and high:
                methods.append(('Ten Cap', float(min(low, high)), float(max(low, high)), '#f59e0b'))
        elif isinstance(tc, (int, float)) and tc > 0:
            methods.append(('Ten Cap', float(tc) * 0.9, float(tc) * 1.1, '#f59e0b'))

        # Equity Bond
        eb = buy_prices.get('equityBondBuyPrice')
        if isinstance(eb, dict) and eb.get('low') and eb.get('high'):
            methods.append(('Equity Bond', float(eb['low']), float(eb['high']), '#8b5cf6'))
        elif isinstance(eb, (int, float)) and eb > 0:
            methods.append(('Equity Bond', float(eb) * 0.9, float(eb) * 1.1, '#8b5cf6'))

        if methods and current_price > 0:
            path = generate_price_range_chart(
                methods, float(current_price),
                title=f'{data.ticker} Full Price Ranges',
            )
            temp_charts.append(path)
            embed_chart(doc, path)
    except Exception:
        pass


def generate_pitch_deck_docx(ticker, base_dir=None):
    """Generate Pitch Deck Word document for the given ticker."""

    data = ReportData(ticker, 'pitch-deck')
    company_name = data.get_company_name()
    verdict = data.get_overall_verdict()

    doc = create_thesis_doc()
    temp_charts = []
    all_citations = []

    # ── Title Page ───────────────────────────────────────────────────────────
    add_title_page(
        doc, ticker, company_name, 'Pitch Deck',
        subtitle='Investment Research Analysis',
        verdict=verdict,
    )

    # Pipeline flow on cover
    try:
        flow_path = generate_pipeline_flow('Pitch Deck')
        temp_charts.append(flow_path)
        embed_chart(doc, flow_path)
    except Exception:
        pass

    # Build a normalized-key index so legacy reports resolve to current keys.
    section_index = {}
    for raw_key in data.get_section_keys():
        norm = normalize_section_key(raw_key)
        s = data.get_section(raw_key)
        if s and norm not in section_index:
            section_index[norm] = (raw_key, s)

    # ── Verdict Scorecard ────────────────────────────────────────────────────
    try:
        scorecard_sections = []
        for norm_key, (raw_key, section) in section_index.items():
            if norm_key == 'investment_verdict':
                continue
            name = section.get('title', norm_key.replace('_', ' ').title())
            v = section.get('verdict', 'N/A')
            conf = section.get('confidence', '')
            signal = str(section.get('verdictRationale', ''))[:60]
            scorecard_sections.append((name, v, conf, signal))

        if scorecard_sections:
            chart_path = generate_verdict_scorecard(
                scorecard_sections, title=f'{ticker} Pitch Deck Scorecard'
            )
            temp_charts.append(chart_path)
            add_section_heading(doc, 'Verdict Scorecard', level=1)
            embed_chart(doc, chart_path)
            add_verdict_table(doc, scorecard_sections)
    except Exception:
        pass

    # ── Per-Section Rendering (7-group structure) ────────────────────────────
    for group_num, group in enumerate(TOP_LEVEL_GROUPS, start=1):
        sub_keys = group['subsection_keys']
        present_subs = [(sk, section_index[sk]) for sk in sub_keys if sk in section_index]
        if not present_subs:
            continue

        add_section_heading(doc, f'{group_num}. {group["title"]}', level=1)

        for norm_key, (raw_key, section) in present_subs:
            sub_title = section.get('title', norm_key.replace('_', ' ').title())
            add_section_heading(doc, sub_title, level=2)

            # Section verdict line
            sec_verdict = section.get('verdict', '')
            if sec_verdict:
                p = doc.add_paragraph()
                run = p.add_run(f'Verdict: {sec_verdict}')
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Arial'
                run.font.color.rgb = VERDICT_COLORS_RGB.get(
                    str(sec_verdict).upper().strip(), SLATE_600
                )
                sec_conf = section.get('confidence', '')
                if sec_conf:
                    run = p.add_run(f'  |  Confidence: {sec_conf}')
                    run.font.size = Pt(10)
                    run.font.name = 'Arial'
                    run.font.color.rgb = SLATE_600

            # Strategic chart for specific sections (keyed by NORMALIZED key)
            chart_type = CHART_SECTIONS.get(norm_key)
            if chart_type == 'revenue':
                _add_revenue_chart(doc, data, temp_charts)
            elif chart_type == 'fcf':
                _add_fcf_chart(doc, data, temp_charts)
                _add_capex_donut(doc, data, temp_charts)
                _add_capital_allocation_stack(doc, data, temp_charts)
            elif chart_type == 'debt_equity':
                _add_debt_equity_chart(doc, data, temp_charts)
            elif chart_type == 'price_range':
                _add_price_range_chart(doc, data, temp_charts)
            elif norm_key == 'market_position':
                _add_peer_comparison(doc, data, temp_charts)
            elif norm_key == 'moat_analysis':
                _add_moat_radar(doc, data, temp_charts)

            # Narrative
            narrative = get_narrative(section)
            if narrative:
                add_body_paragraphs(doc, narrative)

            # Tables
            tables = get_tables(section)
            for t in tables:
                if t.get('title'):
                    add_section_heading(doc, t['title'], level=3)
                add_styled_table(doc, t.get('headers', []), t.get('rows', []))

            # Red flags (skip on accounting_red_flags — categories block covers)
            if norm_key != 'accounting_red_flags':
                flags = get_red_flags(section)
                if flags:
                    render_red_flags(doc, flags)

            # §4d Accounting Red Flags categories block
            if norm_key == 'accounting_red_flags':
                render_accounting_red_flags(doc, section)

            # Verdict box after each subsection narrative
            render_verdict_box(doc, section)

            # Investment verdict gets the closing pre-decision check block
            if norm_key == 'investment_verdict':
                render_pre_decision_check(doc, section)

            # Collect citations for end-of-doc section
            cites = get_citations(section)
            all_citations.extend(cites)

    # ── Citations ────────────────────────────────────────────────────────────
    if all_citations:
        add_citations_section(doc, all_citations)

    # ── Save ─────────────────────────────────────────────────────────────────
    output_path = os.path.join(data.report_dir, 'pitch-deck.docx')
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)

    cleanup_temp_charts(temp_charts)
    print(f'Pitch Deck Word doc saved: {output_path}')
    return output_path


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('Usage: python3 scripts/pdf/generate_pitch_deck_docx.py TICKER')
        sys.exit(1)
    try:
        ticker = normalize_ticker(sys.argv[1])
    except ValueError as error:
        print(error)
        sys.exit(1)
    generate_pitch_deck_docx(ticker)
