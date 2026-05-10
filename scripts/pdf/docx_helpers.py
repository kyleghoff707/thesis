#!/usr/bin/env python3
"""
Shared helpers for Thesis-branded Word document generation.

Provides reusable building blocks for all 3 Word doc generators:
- Document creation with Thesis font/color styles
- Title pages with branding
- Styled tables with teal headers and alternating row shading
- Verdict tables with color-coded cells
- Checklist tables for Final Thesis
- Chart image embedding
- Red flag rendering
- Citations section
- Temporary chart cleanup

Usage:
    from scripts.pdf.docx_helpers import create_thesis_doc, add_title_page, add_styled_table
"""

import os
import re

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


# ── Thesis Colors as RGBColor objects ────────────────────────────────────────

TEAL_500 = RGBColor(15, 118, 110)
TEAL_100 = RGBColor(204, 251, 241)
TEAL_50 = RGBColor(240, 253, 250)
SLATE_800 = RGBColor(30, 41, 59)
SLATE_600 = RGBColor(71, 85, 105)
SLATE_200 = RGBColor(226, 232, 240)
RED_500 = RGBColor(239, 68, 68)
AMBER_500 = RGBColor(245, 158, 11)
GREEN_500 = RGBColor(34, 197, 94)
WHITE = RGBColor(255, 255, 255)

# ── Hex strings for cell shading XML ─────────────────────────────────────────

TEAL_500_HEX = '0F766E'
TEAL_50_HEX = 'F0FDFA'
RED_500_HEX = 'EF4444'
AMBER_500_HEX = 'F59E0B'
GREEN_500_HEX = '22C55E'
SLATE_100_HEX = 'F1F5F9'

VERDICT_COLORS_HEX = {
    'PASS': GREEN_500_HEX,
    'FAIL': RED_500_HEX,
    'PARTIAL': AMBER_500_HEX,
    'WATCHLIST': AMBER_500_HEX,
    'CONTEXT': AMBER_500_HEX,
}

VERDICT_COLORS_RGB = {
    'PASS': GREEN_500,
    'FAIL': RED_500,
    'PARTIAL': AMBER_500,
    'WATCHLIST': AMBER_500,
    'CONTEXT': AMBER_500,
}


def create_thesis_doc():
    """
    Create a Word Document with Thesis branding applied to default styles.

    Sets Arial font, Thesis color palette for headings and body text,
    and consistent paragraph spacing.

    Returns:
        Document: A configured python-docx Document instance
    """
    doc = Document()

    # Normal style: Arial, 10pt, slate-800
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    font.color.rgb = SLATE_800
    pf = style.paragraph_format
    pf.space_after = Pt(6)

    # Heading 1: Arial, 14pt, Bold, Teal-500
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Arial'
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.font.color.rgb = TEAL_500
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)

    # Heading 2: Arial, 12pt, Bold, Slate-800
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Arial'
    h2.font.size = Pt(12)
    h2.font.bold = True
    h2.font.color.rgb = SLATE_800
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(4)

    # Heading 3: Arial, 10pt, Bold, Slate-600
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Arial'
    h3.font.size = Pt(10)
    h3.font.bold = True
    h3.font.color.rgb = SLATE_600
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)

    return doc


def add_title_page(doc, ticker, company_name, stage_title, subtitle='', verdict=''):
    """
    Add a branded title page with company info, stage title, and optional verdict.

    Args:
        doc: Document instance
        ticker: Stock ticker symbol
        company_name: Full company name
        stage_title: E.g., 'One Pager', 'Pitch Deck', 'Final Thesis'
        subtitle: Optional subtitle text
        verdict: Optional verdict string (PASS/FAIL/PARTIAL/WATCHLIST)
    """
    # Spacer
    for _ in range(3):
        doc.add_paragraph('')

    # Company name — large centered teal
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(company_name)
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = TEAL_500
    run.font.name = 'Arial'

    # Ticker
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'({ticker})')
    run.font.size = Pt(14)
    run.font.color.rgb = SLATE_600
    run.font.name = 'Arial'

    # Stage title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(stage_title)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = SLATE_800
    run.font.name = 'Arial'

    # Subtitle
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subtitle)
        run.font.size = Pt(11)
        run.font.italic = True
        run.font.color.rgb = SLATE_600
        run.font.name = 'Arial'

    # Verdict
    if verdict:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        v_upper = str(verdict).upper().strip()
        run = p.add_run(f'Overall Verdict: {v_upper}')
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.name = 'Arial'
        run.font.color.rgb = VERDICT_COLORS_RGB.get(v_upper, SLATE_600)

    # Spacer
    doc.add_paragraph('')

    # Generated by footer
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Generated by Thesis')
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = SLATE_600
    run.font.name = 'Arial'

    # Page break
    doc.add_page_break()


def _set_cell_shading(cell, hex_color):
    """Apply background shading to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_font(cell, size_pt=9, bold=False, color=None, name='Arial'):
    """Set font properties on all runs in a cell paragraph."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(size_pt)
            run.font.bold = bold
            run.font.name = name
            if color:
                run.font.color.rgb = color


def add_styled_table(doc, headers, rows, col_widths=None):
    """
    Create a table with teal header row and alternating row shading.

    Args:
        doc: Document instance
        headers: List of header strings
        rows: List of row lists (each inner list = one row of cell values)
        col_widths: Optional list of Inches values for column widths

    Returns:
        Table: The created table object
    """
    if not headers:
        return None

    n_rows = len(rows) if rows else 0
    table = doc.add_table(rows=1 + n_rows, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(header)
        _set_cell_shading(cell, TEAL_500_HEX)
        _set_cell_font(cell, size_pt=9, bold=True, color=WHITE)

    # Data rows
    for row_idx, row_data in enumerate(rows or []):
        for col_idx in range(len(headers)):
            cell = table.rows[row_idx + 1].cells[col_idx]
            val = str(row_data[col_idx]) if col_idx < len(row_data) else ''
            cell.text = val
            _set_cell_font(cell, size_pt=9, color=SLATE_800)
            # Alternating row shading
            if row_idx % 2 == 0:
                _set_cell_shading(cell, TEAL_50_HEX)

    # Apply column widths if provided
    if col_widths and len(col_widths) == len(headers):
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = width

    # Add spacing after table
    doc.add_paragraph('')

    return table


def add_verdict_table(doc, sections):
    """
    Create a verdict scorecard table with color-coded verdict cells.

    Args:
        doc: Document instance
        sections: List of (name, verdict, confidence, signal) tuples
    """
    if not sections:
        return

    headers = ['Section', 'Verdict', 'Confidence', 'Signal']
    table = doc.add_table(rows=1 + len(sections), cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        _set_cell_shading(cell, TEAL_500_HEX)
        _set_cell_font(cell, size_pt=9, bold=True, color=WHITE)

    # Data rows
    for row_idx, (name, verdict, confidence, signal) in enumerate(sections):
        row = table.rows[row_idx + 1]

        # Section name
        row.cells[0].text = str(name)
        _set_cell_font(row.cells[0], size_pt=9, color=SLATE_800)

        # Verdict cell with color-coded background
        v_str = str(verdict).upper().strip()
        row.cells[1].text = v_str
        v_hex = VERDICT_COLORS_HEX.get(v_str, SLATE_100_HEX)
        _set_cell_shading(row.cells[1], v_hex)
        # White text on colored backgrounds
        v_text_color = WHITE if v_hex != SLATE_100_HEX else SLATE_800
        _set_cell_font(row.cells[1], size_pt=9, bold=True, color=v_text_color)

        # Confidence
        row.cells[2].text = str(confidence)
        _set_cell_font(row.cells[2], size_pt=9, color=SLATE_600)

        # Signal
        row.cells[3].text = str(signal)
        _set_cell_font(row.cells[3], size_pt=9, color=SLATE_600)

        # Alternating row background (only on non-verdict cells)
        if row_idx % 2 == 0:
            _set_cell_shading(row.cells[0], TEAL_50_HEX)
            _set_cell_shading(row.cells[2], TEAL_50_HEX)
            _set_cell_shading(row.cells[3], TEAL_50_HEX)

    doc.add_paragraph('')


def add_section_heading(doc, title, level=1):
    """
    Add a heading at the specified level.

    Args:
        doc: Document instance
        title: Heading text
        level: 1, 2, or 3

    Returns:
        Paragraph: The heading paragraph
    """
    return doc.add_heading(title, level=level)


def add_body_paragraphs(doc, text):
    """
    Split text on double newlines into paragraphs with bold formatting support.

    Handles **bold** markdown syntax by creating bold runs.
    Strips inline <cite> tags and replaces "DataPacket" with "Thesis toolbox".
    Skips empty paragraphs.

    Args:
        doc: Document instance
        text: Multi-paragraph text string
    """
    if not text or not isinstance(text, str):
        return

    # Clean cite tags and internal jargon
    from section_renderers import _clean_narrative
    text = _clean_narrative(text)

    paragraphs = text.split('\n\n')
    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue

        # Standalone bold paragraphs are sub-headers (e.g. "**Capital Allocation**")
        if para_text.startswith('**') and para_text.endswith('**') and len(para_text) < 120:
            add_section_heading(doc, para_text[2:-2], level=3)
            continue

        p = doc.add_paragraph()

        # Parse markdown bold (**text**) into runs
        parts = re.split(r'(\*\*[^*]+\*\*)', para_text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Bold text
                run = p.add_run(part[2:-2])
                run.bold = True
                run.font.name = 'Arial'
                run.font.size = Pt(10)
                run.font.color.rgb = SLATE_800
            else:
                if part:
                    run = p.add_run(part)
                    run.font.name = 'Arial'
                    run.font.size = Pt(10)
                    run.font.color.rgb = SLATE_800


def embed_chart(doc, chart_path, width_inches=5.5):
    """
    Embed a PNG chart image in the document.

    Args:
        doc: Document instance
        chart_path: Path to the PNG file
        width_inches: Image width in inches (default 5.5 fits within margins)
    """
    if not chart_path or not os.path.exists(chart_path):
        return

    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(chart_path, width=Inches(width_inches))
        # Add small spacing after
        doc.add_paragraph('')
    except Exception:
        # Skip silently if image can't be embedded
        pass


def add_red_flags(doc, flags):
    """
    Render red flags as a bulleted list with red triangle prefix.

    Args:
        doc: Document instance
        flags: List of flag strings
    """
    if not flags:
        return

    doc.add_heading('Red Flags', level=3)

    for flag in flags:
        p = doc.add_paragraph()
        # Red triangle prefix
        run = p.add_run('\u26A0 ')
        run.font.color.rgb = RED_500
        run.font.size = Pt(10)
        run.font.name = 'Arial'
        # Flag text
        run = p.add_run(str(flag))
        run.font.size = Pt(10)
        run.font.name = 'Arial'
        run.font.color.rgb = SLATE_800


def add_citations_section(doc, all_citations):
    """
    Add a numbered citations/sources section.

    Args:
        doc: Document instance
        all_citations: List of {ref, text, source} dicts
    """
    if not all_citations:
        return

    doc.add_heading('Citations & Sources', level=1)

    for i, cite in enumerate(all_citations, 1):
        if not isinstance(cite, dict):
            continue
        ref = cite.get('ref', '')
        text = cite.get('text', '')
        source = cite.get('source', '')

        p = doc.add_paragraph()
        # Number
        run = p.add_run(f'[{i}] ')
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = TEAL_500
        run.font.name = 'Arial'

        # Ref text
        if ref:
            run = p.add_run(f'{ref}')
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.color.rgb = SLATE_800
            run.font.name = 'Arial'

        if text:
            run = p.add_run(f' = {text}')
            run.font.size = Pt(9)
            run.font.color.rgb = SLATE_600
            run.font.name = 'Arial'

        if source:
            run = p.add_run(f' ({source})')
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = SLATE_600
            run.font.name = 'Arial'


def add_checklist_table(doc, items):
    """
    Create a checklist table for Final Thesis sections with color-coded verdict cells.

    Args:
        doc: Document instance
        items: List of {number, item, verdict, confidence, evidence} dicts
    """
    if not items:
        return

    headers = ['#', 'Item', 'Verdict', 'Confidence']
    table = doc.add_table(rows=1 + len(items), cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        _set_cell_shading(cell, TEAL_500_HEX)
        _set_cell_font(cell, size_pt=9, bold=True, color=WHITE)

    # Data rows
    for row_idx, item in enumerate(items):
        if not isinstance(item, dict):
            continue

        row = table.rows[row_idx + 1]

        # Number
        row.cells[0].text = str(item.get('number', row_idx + 1))
        _set_cell_font(row.cells[0], size_pt=9, color=SLATE_800)

        # Item text
        item_text = str(item.get('item', ''))
        row.cells[1].text = item_text
        _set_cell_font(row.cells[1], size_pt=9, color=SLATE_800)

        # Verdict with color-coded background
        verdict = str(item.get('verdict', 'N/A')).upper().strip()
        row.cells[2].text = verdict
        v_hex = VERDICT_COLORS_HEX.get(verdict, SLATE_100_HEX)
        _set_cell_shading(row.cells[2], v_hex)
        v_text_color = WHITE if v_hex != SLATE_100_HEX else SLATE_800
        _set_cell_font(row.cells[2], size_pt=9, bold=True, color=v_text_color)

        # Confidence
        row.cells[3].text = str(item.get('confidence', ''))
        _set_cell_font(row.cells[3], size_pt=9, color=SLATE_600)

        # Alternating row shading on non-verdict cells
        if row_idx % 2 == 0:
            _set_cell_shading(row.cells[0], TEAL_50_HEX)
            _set_cell_shading(row.cells[1], TEAL_50_HEX)
            _set_cell_shading(row.cells[3], TEAL_50_HEX)

    doc.add_paragraph('')

    # Evidence detail paragraphs for items with substantial evidence
    evidence_items = [item for item in items
                      if isinstance(item, dict)
                      and isinstance(item.get('evidence', ''), str)
                      and len(item.get('evidence', '')) > 100]

    if evidence_items:
        doc.add_heading('Evidence Details', level=3)
        for item in evidence_items:
            num = item.get('number', '?')
            item_name = str(item.get('item', ''))[:80]
            p = doc.add_paragraph()
            run = p.add_run(f'#{num}: {item_name}')
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = TEAL_500
            run.font.name = 'Arial'

            p = doc.add_paragraph()
            run = p.add_run(str(item.get('evidence', '')))
            run.font.size = Pt(9)
            run.font.color.rgb = SLATE_600
            run.font.name = 'Arial'


def _section_data_dict(section):
    """Return section['data'] as a dict, parsing JSON-string payloads safely."""
    if not isinstance(section, dict):
        return {}
    data = section.get('data', {})
    if isinstance(data, str):
        import json
        try:
            data = json.loads(data)
        except (json.JSONDecodeError, ValueError):
            return {}
    return data if isinstance(data, dict) else {}


def _camel_to_words_docx(s):
    """Convert camelCase or snake_case to 'Title Words' for DOCX labels."""
    s = str(s).replace('_', ' ')
    s = re.sub(r'([a-z])([A-Z])', r'\1 \2', s)
    return s[:1].upper() + s[1:] if s else s


def render_verdict_box_docx(doc, section):
    """
    Render the verdict callout for a Final Thesis prose section as a
    1-cell shaded table (DOCX equivalent of the bordered PDF box).

    Reads `data.verdict`. Renders nothing if missing or not a dict —
    graceful for legacy reports.
    """
    data = _section_data_dict(section)
    verdict = data.get('verdict')
    if not isinstance(verdict, dict):
        return

    overall = str(verdict.get('overall', 'WATCHLIST')).upper().strip()
    color_map = {
        'PASS': (GREEN_500, GREEN_500_HEX),
        'WATCHLIST': (AMBER_500, AMBER_500_HEX),
        'PARTIAL': (AMBER_500, AMBER_500_HEX),
        'CONTEXT': (AMBER_500, AMBER_500_HEX),
        'FAIL': (RED_500, RED_500_HEX),
    }
    rgb, _ = color_map.get(overall, (SLATE_600, SLATE_100_HEX))

    # 1-cell single-row table with light shading + coloured text
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    _set_cell_shading(cell, SLATE_100_HEX)

    # Heading
    title = section.get('title') or section.get('key', 'Section')
    p = cell.paragraphs[0]
    run = p.add_run(f'{title} verdict')
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    run.font.color.rgb = rgb

    # Verdict-detail lines (skip 'overall')
    for k, v in verdict.items():
        if k == 'overall':
            continue
        label = _camel_to_words_docx(k)
        p = cell.add_paragraph()
        run = p.add_run(f'{label}: {v}')
        run.font.size = Pt(10)
        run.font.name = 'Arial'
        run.font.color.rgb = SLATE_600

    # Overall stamp
    p = cell.add_paragraph()
    run = p.add_run(f'Verdict: {overall}')
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    run.font.color.rgb = rgb

    doc.add_paragraph('')


def render_promise_tracker_docx(doc, section):
    """
    Render the Management Promise Tracker as a 5-column DOCX subsection.

    Pulled from `data.promises`. Renders nothing if absent.
    """
    data = _section_data_dict(section)
    promises = data.get('promises', [])
    if not isinstance(promises, list) or not promises:
        return

    doc.add_heading('Management Promise Tracker', level=2)
    headers = ['Quarter', 'Category', 'Promise', 'Evidence', 'Status']
    rows = []
    for p in promises:
        if not isinstance(p, dict):
            continue
        rows.append([
            str(p.get('quarterYear', p.get('quarter', ''))),
            str(p.get('category', '')),
            str(p.get('quote', p.get('promise', '')))[:160],
            str(p.get('evidence', ''))[:160],
            str(p.get('status', '')),
        ])
    if rows:
        add_styled_table(doc, headers, rows)


def render_trade_plan_docx(doc, section):
    """
    Render Section 7 Trade Plan: position sizing, entry tranches table,
    sell rules list, PACE plan, forcing question.
    """
    data = _section_data_dict(section)

    # Position sizing
    sizing = data.get('positionSizing')
    if sizing:
        doc.add_heading('Position Sizing', level=2)
        if isinstance(sizing, str):
            add_body_paragraphs(doc, sizing)
        elif isinstance(sizing, dict):
            for label, value in sizing.items():
                if value is None or value == '':
                    continue
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(f'{_camel_to_words_docx(label)}: {value}')
                run.font.size = Pt(10)
                run.font.name = 'Arial'
                run.font.color.rgb = SLATE_800

    # Entry tranches table
    tranches = data.get('tranches', data.get('entryTranches', []))
    if isinstance(tranches, list) and tranches:
        doc.add_heading('Entry Tranches', level=2)
        headers = ['Tranche', 'Size', 'Trigger Price', 'Rationale']
        rows = []
        for t in tranches:
            if not isinstance(t, dict):
                continue
            rows.append([
                str(t.get('tranche', t.get('label', ''))),
                str(t.get('size', t.get('sizePct', ''))),
                str(t.get('triggerPrice', t.get('trigger', ''))),
                str(t.get('rationale', ''))[:160],
            ])
        if rows:
            add_styled_table(doc, headers, rows)

    # Sell rules
    sell_rules = data.get('sellRules', [])
    if isinstance(sell_rules, list) and sell_rules:
        doc.add_heading('Sell Rules', level=2)
        for r in sell_rules:
            if isinstance(r, dict):
                trigger = r.get('trigger', '')
                action = r.get('action', '')
                threshold = r.get('threshold', '')
                line = f'{trigger}: {action}' if action else str(trigger)
                if threshold:
                    line += f' (threshold: {threshold})'
            elif isinstance(r, str):
                line = r
            else:
                continue
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(line)
            run.font.size = Pt(10)
            run.font.name = 'Arial'
            run.font.color.rgb = SLATE_800

    # PACE plan
    pace = data.get('pacePlan')
    if isinstance(pace, dict):
        doc.add_heading('PACE Plan', level=2)
        for label in ('primary', 'alternative', 'contingency', 'emergency'):
            value = pace.get(label, '')
            if value:
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(f'{label.capitalize()}: {value}')
                run.font.size = Pt(10)
                run.font.name = 'Arial'
                run.font.color.rgb = SLATE_800

    # Forcing question
    fq = data.get('forcingQuestion')
    if fq:
        doc.add_heading('Forcing Question', level=2)
        p = doc.add_paragraph()
        run = p.add_run(str(fq))
        run.font.italic = True
        run.font.size = Pt(10)
        run.font.name = 'Arial'
        run.font.color.rgb = SLATE_600


def render_watchpoints_docx(doc, section):
    """
    Render the "What we're monitoring" subsection at the end of §6 Compose.

    Reads `data.watchpoints`. Renders nothing if absent.
    """
    data = _section_data_dict(section)
    watchpoints = data.get('watchpoints', [])
    if not isinstance(watchpoints, list) or not watchpoints:
        return

    doc.add_heading("What we're monitoring", level=2)
    for wp in watchpoints:
        if not isinstance(wp, dict):
            continue
        metric = wp.get('metric', '')
        current = wp.get('currentValue', wp.get('current', ''))
        threshold = wp.get('threshold', '')
        direction = str(wp.get('direction', '')).lower()
        if direction == 'below':
            change = 'drops below'
        elif direction == 'above':
            change = 'rises above'
        else:
            change = 'crosses'
        line = f'{metric}.'
        if current != '' and current is not None:
            line += f' Currently {current}.'
        if threshold != '' and threshold is not None:
            line += f' Re-evaluate if it {change} {threshold}.'
        src = wp.get('sourceInversionId')
        if src is not None and src != '':
            line += f' (Source: bear inversion #{src}.)'
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(line)
        run.font.size = Pt(10)
        run.font.name = 'Arial'
        run.font.color.rgb = SLATE_800


# =============================================================================
# PITCH DECK REDESIGN (2026-05-09) — VERDICT BOX + REDESIGN-SPECIFIC HELPERS
# =============================================================================

# Verdict colors matching the redesign-plan spec.
VERDICT_BOX_HEX = {
    'PASS':       '4CAF50',
    'WATCHLIST':  'FF9800',
    'PARTIAL':    'FF9800',
    'CONTEXT':    'FF9800',
    'FAIL':       'F44336',
}


def _verdict_box_hex(verdict):
    v = str(verdict or '').upper().strip()
    return VERDICT_BOX_HEX.get(v, '64748B')  # slate fallback


def _verdict_rgb(verdict):
    """Return RGBColor matching the verdict-box hex palette."""
    h = _verdict_box_hex(verdict)
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _strip_cite_tags(text):
    """Lightweight wrapper that strips <cite> tags. Local import keeps this
    module free of a hard dependency on section_renderers at import time."""
    if not text or not isinstance(text, str):
        return text or ''
    try:
        from section_renderers import _clean_narrative
        return _clean_narrative(text)
    except Exception:
        return text


def render_verdict_box(doc, section):
    """Render a color-coded bordered call-out summarizing the section verdict.

    DOCX-equivalent of the PDF verdict box — a single-cell bordered table
    with a colored left edge stripe (via cell-border XML), a colored heading
    row, and body text. Reads section.verdict (PASS/WATCHLIST/FAIL),
    section.verdictRationale, section.confidence, and any section.data.verdictBox
    payload (bullCase/bearCase/whatWouldChange/watchItems).
    """
    if not isinstance(section, dict):
        return
    verdict = str(section.get('verdict', '') or '').upper().strip()
    if not verdict:
        return

    rationale = _strip_cite_tags(section.get('verdictRationale', '') or '')
    confidence = section.get('confidence', '') or ''
    data = section.get('data', {}) if isinstance(section.get('data'), dict) else {}
    box_payload = data.get('verdictBox') if isinstance(data.get('verdictBox'), dict) else {}

    color_hex = _verdict_box_hex(verdict)
    color_rgb = _verdict_rgb(verdict)

    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    _set_cell_shading(cell, 'F8FAFC')

    # Header paragraph
    header_p = cell.paragraphs[0]
    run = header_p.add_run(f'Verdict: {verdict}')
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    run.font.color.rgb = color_rgb
    if confidence:
        run2 = header_p.add_run(f'   |   Confidence: {confidence}')
        run2.font.size = Pt(10)
        run2.font.name = 'Arial'
        run2.font.color.rgb = SLATE_600

    # Rationale
    if rationale:
        p = cell.add_paragraph()
        run = p.add_run(rationale)
        run.font.size = Pt(9)
        run.font.name = 'Arial'
        run.font.color.rgb = SLATE_800

    # Optional structured fields
    if box_payload:
        for label, key in [
            ('Bull Case', 'bullCase'),
            ('Bear Case', 'bearCase'),
            ('What Would Change This', 'whatWouldChange'),
            ('Watch Items', 'watchItems'),
        ]:
            val = box_payload.get(key)
            if not val:
                continue
            p = cell.add_paragraph()
            run = p.add_run(f'{label}: ')
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.name = 'Arial'
            run.font.color.rgb = SLATE_800
            if isinstance(val, list):
                content = '; '.join(str(x) for x in val)
            else:
                content = str(val)
            run = p.add_run(content)
            run.font.size = Pt(9)
            run.font.name = 'Arial'
            run.font.color.rgb = SLATE_600

    # Colored left-edge accent + colored borders via cell-border XML
    tc_pr = cell._tc.get_or_add_tcPr()
    borders_xml = (
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>'
        f'<w:left w:val="single" w:sz="24" w:space="0" w:color="{color_hex}"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>'
        f'</w:tcBorders>'
    )
    tc_pr.append(parse_xml(borders_xml))

    doc.add_paragraph('')


def render_accounting_red_flags(doc, section):
    """Render the §4d Accounting Red Flags categories block.

    Iterates `data.categories[]`, color-coding each by status
    ("Clean" → green, anything else → amber/red). Each category renders
    its `flagsFound[]` list (or a "Clean" note).
    """
    if not isinstance(section, dict):
        return
    data = section.get('data', {})
    if not isinstance(data, dict):
        return
    categories = data.get('categories', [])
    if not isinstance(categories, list) or not categories:
        return

    add_section_heading(doc, 'Accounting Red Flag Categories', level=2)

    for cat in categories:
        if not isinstance(cat, dict):
            continue
        name = str(cat.get('name', cat.get('category', 'Category')))
        status = str(cat.get('status', cat.get('verdict', '')) or '').strip()
        flags = cat.get('flagsFound', cat.get('flags', []))
        if not isinstance(flags, list):
            flags = []

        status_upper = status.upper()
        if status_upper in ('CLEAN', 'PASS', 'OK'):
            color_rgb = _verdict_rgb('PASS')
            display_status = 'Clean'
        elif status_upper in ('FAIL', 'CRITICAL'):
            color_rgb = _verdict_rgb('FAIL')
            display_status = status or 'Issues Found'
        else:
            color_rgb = _verdict_rgb('WATCHLIST')
            display_status = status or 'Issues Found'

        # Category header
        p = doc.add_paragraph()
        run = p.add_run(f'{name}: {display_status}')
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Arial'
        run.font.color.rgb = color_rgb

        # Flags or "Clean" note
        if flags:
            for f in flags:
                if isinstance(f, dict):
                    txt = f.get('flag') or f.get('description') or f.get('detail') or str(f)
                else:
                    txt = str(f)
                p = doc.add_paragraph()
                run = p.add_run('• ')
                run.font.color.rgb = color_rgb
                run.font.size = Pt(10)
                run.font.name = 'Arial'
                run = p.add_run(str(txt))
                run.font.size = Pt(10)
                run.font.name = 'Arial'
                run.font.color.rgb = SLATE_800
        else:
            p = doc.add_paragraph()
            run = p.add_run('No issues identified in this category.')
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.name = 'Arial'
            run.font.color.rgb = SLATE_600


def render_pre_decision_check(doc, section):
    """Render the closing Pre-Decision Quality Check block on Investment Verdict.

    Reads section.data.preDecisionCheck. Visually distinct: bordered single-cell
    table with light-gray fill, italic body text, and bold field labels.
    """
    if not isinstance(section, dict):
        return
    data = section.get('data', {})
    if not isinstance(data, dict):
        return
    pdc = data.get('preDecisionCheck')
    if not pdc:
        return

    confidence_calibration = ''
    anticipated_regret = ''
    free_text = ''
    if isinstance(pdc, dict):
        confidence_calibration = str(pdc.get('confidenceCalibration', '') or '')
        anticipated_regret = str(pdc.get('anticipatedRegret', '') or '')
        free_text = str(pdc.get('text', '') or pdc.get('summary', '') or '')
    elif isinstance(pdc, str):
        free_text = pdc

    if not (confidence_calibration or anticipated_regret or free_text):
        return

    add_section_heading(doc, 'Pre-Decision Quality Check', level=3)

    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    _set_cell_shading(cell, SLATE_100_HEX)

    # First-paragraph anchor (cell starts with one empty paragraph)
    first_used = False
    if confidence_calibration:
        p = cell.paragraphs[0]
        first_used = True
        run = p.add_run('Confidence calibration: ')
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Arial'
        run.font.color.rgb = SLATE_800
        run = p.add_run(confidence_calibration)
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.name = 'Arial'
        run.font.color.rgb = SLATE_600
    if anticipated_regret:
        target = cell.add_paragraph() if first_used else cell.paragraphs[0]
        first_used = True
        run = target.add_run('Anticipated regret: ')
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Arial'
        run.font.color.rgb = SLATE_800
        run = target.add_run(anticipated_regret)
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.name = 'Arial'
        run.font.color.rgb = SLATE_600
    if free_text and not first_used:
        p = cell.paragraphs[0]
        run = p.add_run(free_text)
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.name = 'Arial'
        run.font.color.rgb = SLATE_600

    doc.add_paragraph('')


def cleanup_temp_charts(chart_paths):
    """
    Delete temporary PNG chart files after they have been embedded.

    Args:
        chart_paths: List of file paths to remove
    """
    if not chart_paths:
        return

    for path in chart_paths:
        if path and os.path.exists(path):
            try:
                os.unlink(path)
            except OSError:
                pass
