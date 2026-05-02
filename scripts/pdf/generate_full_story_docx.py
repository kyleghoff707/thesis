#!/usr/bin/env python3
"""
Full Story Word Document Generator — Thes1s-branded .docx export.

Generates a professional Word document from any ticker's Full Story pipeline output
with embedded chart images (checklist summaries, price ranges), checklist tables with
color-coded verdict cells, adversarial debate rendering, and Thes1s branding.

Usage:
    python3 scripts/pdf/generate_full_story_docx.py MNST
"""

import os
import sys

# Ensure scripts/pdf is importable
sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', '..'))

from scripts.pdf.report_data_reader import ReportData
from scripts.pdf.section_renderers import (
    get_narrative, get_tables, get_red_flags, get_citations,
    get_checklist_items,
)
from scripts.pdf.chart_image_generator import (
    generate_checklist_summary, generate_price_range_chart,
)
from scripts.pdf.docx_helpers import (
    create_thes1s_doc, add_title_page, add_styled_table,
    add_section_heading, add_body_paragraphs,
    embed_chart, add_red_flags as render_red_flags,
    add_checklist_table, add_citations_section,
    cleanup_temp_charts, VERDICT_COLORS_RGB, VERDICT_COLORS_HEX,
    TEAL_500, SLATE_600, SLATE_800, WHITE,
)
from docx.shared import Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


# Checklist section keys
CHECKLIST_SECTIONS = {'meaning_checklist', 'moat_checklist', 'management_checklist'}


def _render_debate(doc, debate_outputs, temp_charts):
    """
    Render the 4-step adversarial debate (Bull -> Bear -> Rebuttal -> Judge).

    Args:
        doc: Document instance
        debate_outputs: Dict with bull, bear, bull_rebuttal, judge keys
        temp_charts: List to track temporary chart PNG paths
    """
    if not debate_outputs or not isinstance(debate_outputs, dict):
        return

    # ── Bull Thesis ──────────────────────────────────────────────────────────
    bull = debate_outputs.get('bull', {})
    bull_content = bull.get('content', {})
    if isinstance(bull_content, dict):
        add_section_heading(doc, 'Bull Thesis', level=2)

        thesis = bull_content.get('overallThesis', '')
        if thesis:
            add_body_paragraphs(doc, thesis)

        points = bull_content.get('thesisPoints', [])
        if isinstance(points, list) and points:
            for i, point in enumerate(points, 1):
                if isinstance(point, dict):
                    p = doc.add_paragraph()
                    run = p.add_run(f'{i}. ')
                    run.font.bold = True
                    run.font.size = Pt(10)
                    run.font.name = 'Arial'
                    run.font.color.rgb = TEAL_500

                    pt_text = point.get('point', '')
                    run = p.add_run(str(pt_text))
                    run.font.size = Pt(10)
                    run.font.name = 'Arial'
                    run.font.color.rgb = SLATE_800

                    evidence = point.get('evidence', '')
                    if evidence:
                        p = doc.add_paragraph()
                        run = p.add_run(str(evidence))
                        run.font.size = Pt(9)
                        run.font.italic = True
                        run.font.name = 'Arial'
                        run.font.color.rgb = SLATE_600

    # ── Bear Inversion ───────────────────────────────────────────────────────
    bear = debate_outputs.get('bear', {})
    bear_content = bear.get('content', {})
    if isinstance(bear_content, dict):
        add_section_heading(doc, 'Bear Inversion', level=2)

        bear_case = bear_content.get('overallBearCase', '')
        if bear_case:
            add_body_paragraphs(doc, bear_case)

        inversions = bear_content.get('inversions', [])
        if isinstance(inversions, list) and inversions:
            for inv in inversions:
                if isinstance(inv, dict):
                    severity = inv.get('severity', 'MEDIUM')
                    target = inv.get('targetPoint', '')
                    counter = inv.get('counterArgument', '')

                    p = doc.add_paragraph()
                    # Severity prefix
                    run = p.add_run(f'[{severity}] ')
                    run.font.bold = True
                    run.font.size = Pt(10)
                    run.font.name = 'Arial'
                    from scripts.pdf.docx_helpers import RED_500, AMBER_500
                    if str(severity).upper() == 'HIGH':
                        run.font.color.rgb = RED_500
                    elif str(severity).upper() == 'CRITICAL':
                        run.font.color.rgb = RED_500
                    else:
                        run.font.color.rgb = AMBER_500

                    if target:
                        run = p.add_run(f'{target}: ')
                        run.font.bold = True
                        run.font.size = Pt(10)
                        run.font.name = 'Arial'
                        run.font.color.rgb = SLATE_800

                    if counter:
                        run = p.add_run(str(counter))
                        run.font.size = Pt(10)
                        run.font.name = 'Arial'
                        run.font.color.rgb = SLATE_800

    # ── Bull Rebuttal ────────────────────────────────────────────────────────
    rebuttal = debate_outputs.get('bull_rebuttal', {})
    rebuttal_content = rebuttal.get('content', {})
    if isinstance(rebuttal_content, dict):
        add_section_heading(doc, 'Bull Rebuttal', level=2)

        rebuttals = rebuttal_content.get('rebuttals', [])
        if isinstance(rebuttals, list) and rebuttals:
            for reb in rebuttals:
                if isinstance(reb, dict):
                    strength = reb.get('rebuttalStrength', 'MEDIUM')
                    bear_pt = reb.get('bearPoint', '')
                    rebuttal_text = reb.get('rebuttal', '')

                    p = doc.add_paragraph()
                    run = p.add_run(f'[{strength}] ')
                    run.font.bold = True
                    run.font.size = Pt(10)
                    run.font.name = 'Arial'
                    from scripts.pdf.docx_helpers import GREEN_500
                    if str(strength).upper() in ('STRONG', 'HIGH'):
                        run.font.color.rgb = GREEN_500
                    else:
                        run.font.color.rgb = AMBER_500

                    if bear_pt:
                        run = p.add_run(f'Re: {bear_pt}: ')
                        run.font.bold = True
                        run.font.size = Pt(10)
                        run.font.name = 'Arial'
                        run.font.color.rgb = SLATE_800

                    if rebuttal_text:
                        run = p.add_run(str(rebuttal_text))
                        run.font.size = Pt(10)
                        run.font.name = 'Arial'
                        run.font.color.rgb = SLATE_800

    # ── Judge Verdict ────────────────────────────────────────────────────────
    judge = debate_outputs.get('judge', {})
    judge_content = judge.get('content', {})
    if isinstance(judge_content, dict):
        add_section_heading(doc, 'Judge Verdict', level=2)

        overall_verdict = judge_content.get('overallVerdict', {})
        if isinstance(overall_verdict, dict):
            direction = overall_verdict.get('direction', '')
            unresolved = overall_verdict.get('unresolvedCount', 0)
            summary = overall_verdict.get('summary', '')
            implication = overall_verdict.get('investmentImplication', '')

            # Direction and unresolved count
            p = doc.add_paragraph()
            run = p.add_run(f'Direction: {direction}')
            run.font.bold = True
            run.font.size = Pt(12)
            run.font.name = 'Arial'
            if str(direction).lower() == 'bull':
                from scripts.pdf.docx_helpers import GREEN_500
                run.font.color.rgb = GREEN_500
            else:
                from scripts.pdf.docx_helpers import RED_500
                run.font.color.rgb = RED_500

            if unresolved:
                run = p.add_run(f'  |  Unresolved: {unresolved}')
                run.font.size = Pt(10)
                run.font.name = 'Arial'
                run.font.color.rgb = SLATE_600

            # Summary
            if summary:
                add_body_paragraphs(doc, summary)

            # Investment implication
            if implication:
                add_section_heading(doc, 'Investment Implication', level=3)
                add_body_paragraphs(doc, implication)

        # Exchange table
        exchanges = judge_content.get('exchanges', [])
        if isinstance(exchanges, list) and exchanges:
            add_section_heading(doc, 'Debate Exchanges', level=3)

            headers = ['Topic', 'Bull', 'Bear', 'Verdict', 'Reasoning']
            rows = []
            for ex in exchanges:
                if isinstance(ex, dict):
                    rows.append([
                        str(ex.get('topic', '')),
                        str(ex.get('bullStrength', '')),
                        str(ex.get('bearStrength', '')),
                        str(ex.get('verdict', '')),
                        str(ex.get('reasoning', ''))[:120],
                    ])

            if rows:
                table = add_styled_table(doc, headers, rows)

                # Color-code the Verdict column (index 3)
                if table:
                    for row_idx in range(len(rows)):
                        cell = table.rows[row_idx + 1].cells[3]
                        verdict_text = rows[row_idx][3].upper().strip()
                        v_hex = VERDICT_COLORS_HEX.get(verdict_text)
                        if not v_hex:
                            # Map Bull/Bear/Unresolved to colors
                            if 'BULL' in verdict_text:
                                v_hex = '22C55E'
                            elif 'BEAR' in verdict_text:
                                v_hex = 'EF4444'
                            elif 'UNRESOLVED' in verdict_text or 'DRAW' in verdict_text:
                                v_hex = 'F59E0B'
                        if v_hex:
                            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{v_hex}"/>')
                            cell._tc.get_or_add_tcPr().append(shading)
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.color.rgb = WHITE


def generate_full_story_docx(ticker, base_dir=None):
    """Generate Full Story Word document for the given ticker."""

    data = ReportData(ticker, 'full-story', base_dir=base_dir)
    company_name = data.get_company_name()

    doc = create_thes1s_doc()
    temp_charts = []
    all_citations = []

    # ── Title Page (no single verdict — the debate IS the verdict) ───────────
    add_title_page(
        doc, ticker, company_name, 'Full Story',
        subtitle='Investment Conviction Analysis',
    )

    # ── Section Rendering ────────────────────────────────────────────────────
    for key in data.get_section_keys():
        section = data.get_section(key)
        title = section.get('title', key.replace('_', ' ').title())

        add_section_heading(doc, title, level=1)

        # Section verdict
        sec_verdict = section.get('verdict', '')
        if sec_verdict:
            p = doc.add_paragraph()
            run = p.add_run(f'Verdict: {sec_verdict}')
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.name = 'Arial'
            run.font.color.rgb = VERDICT_COLORS_RGB.get(
                str(sec_verdict).upper().strip(), SLATE_600
            )

        # ── Checklist Sections ───────────────────────────────────────────────
        if key in CHECKLIST_SECTIONS:
            items = get_checklist_items(section)
            if items:
                # Checklist summary chart
                try:
                    chart_path = generate_checklist_summary(
                        items, title=f'{title} Results'
                    )
                    temp_charts.append(chart_path)
                    embed_chart(doc, chart_path)
                except Exception:
                    pass

                # Checklist table with color-coded verdicts
                add_checklist_table(doc, items)

        # ── Valuation Confirmation — Price Range Chart ───────────────────────
        if key == 'valuation_confirmation':
            try:
                buy_prices = data.get_buy_prices()
                current_price = buy_prices.get('currentPrice', 0) or data.get_current_price()
                if current_price:
                    methods = []
                    mos = buy_prices.get('mosBuyPrice')
                    if isinstance(mos, dict) and mos.get('low') and mos.get('high'):
                        methods.append(('MOS', float(mos['low']), float(mos['high']), '#22c55e'))
                    elif isinstance(mos, (int, float)) and mos > 0:
                        methods.append(('MOS', float(mos) * 0.9, float(mos) * 1.1, '#22c55e'))

                    pbt = buy_prices.get('pbtBuyPrice')
                    if isinstance(pbt, dict) and pbt.get('low') and pbt.get('high'):
                        methods.append(('PBT', float(pbt['low']), float(pbt['high']), '#3b82f6'))
                    elif isinstance(pbt, (int, float)) and pbt > 0:
                        methods.append(('PBT', float(pbt) * 0.9, float(pbt) * 1.1, '#3b82f6'))

                    tc = buy_prices.get('tenCapPrice')
                    if isinstance(tc, dict):
                        low = tc.get('ruleOne') or tc.get('low')
                        high = tc.get('graham') or tc.get('high')
                        if low and high:
                            methods.append(('Ten Cap', float(min(low, high)), float(max(low, high)), '#f59e0b'))
                    elif isinstance(tc, (int, float)) and tc > 0:
                        methods.append(('Ten Cap', float(tc) * 0.9, float(tc) * 1.1, '#f59e0b'))

                    eb = buy_prices.get('equityBondBuyPrice')
                    if isinstance(eb, dict) and eb.get('low') and eb.get('high'):
                        methods.append(('Equity Bond', float(eb['low']), float(eb['high']), '#8b5cf6'))
                    elif isinstance(eb, (int, float)) and eb > 0:
                        methods.append(('Equity Bond', float(eb) * 0.9, float(eb) * 1.1, '#8b5cf6'))

                    if methods:
                        path = generate_price_range_chart(
                            methods, float(current_price),
                            title=f'{ticker} Buy Price Ranges',
                        )
                        temp_charts.append(path)
                        embed_chart(doc, path)
            except Exception:
                pass

        # Narrative
        narrative = get_narrative(section)
        if narrative:
            add_body_paragraphs(doc, narrative)

        # Tables
        tables = get_tables(section)
        for t in tables:
            if t.get('title'):
                add_section_heading(doc, t['title'], level=3)
            add_styled_table(doc, t.get('headers', []), t.get('rows', []))

        # Red flags
        flags = get_red_flags(section)
        if flags:
            render_red_flags(doc, flags)

        # Collect citations
        cites = get_citations(section)
        all_citations.extend(cites)

    # ── Inversion & Rebuttal — Debate Rendering ─────────────────────────────
    debate_outputs = data.get_debate_outputs()
    if debate_outputs:
        add_section_heading(doc, 'Adversarial Debate', level=1)
        _render_debate(doc, debate_outputs, temp_charts)

    # ── Citations ────────────────────────────────────────────────────────────
    if all_citations:
        add_citations_section(doc, all_citations)

    # ── Save ─────────────────────────────────────────────────────────────────
    output_path = os.path.join(data.report_dir, 'full-story.docx')
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)

    cleanup_temp_charts(temp_charts)
    print(f'Full Story Word doc saved: {output_path}')
    return output_path


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('Usage: python3 scripts/pdf/generate_full_story_docx.py TICKER')
        sys.exit(1)
    generate_full_story_docx(sys.argv[1])
